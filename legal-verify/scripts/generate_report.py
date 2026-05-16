#!/usr/bin/env python3
"""
法条核验报告 DOCX 生成器 (generate_report.py)

将 JSON 格式的核验结果转换为专业排版的 DOCX 报告。

用法：
    python generate_report.py <JSON文件路径> [--output <输出路径>] [--title <报告标题>]

JSON schema 详见 SKILL.md 中的 "核验结果 JSON Schema" 章节。
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path


# ============================================================
# 颜色与样式常量
# ============================================================

# 主题色
DARK_BLUE = "1B3A5C"
RED = "C0392B"
GREEN = "27AE60"
ORANGE = "E67E22"
GOLD = "C49B3A"
LIGHT_GRAY = "F2F2F2"
MEDIUM_GRAY = "95A5A6"
DARK_GRAY = "2C3E50"
WHITE = "FFFFFF"

# 页面尺寸常量（A4 标准页边距下可用内容宽度约 15.5cm）
PAGE_CONTENT_WIDTH = 15.5  # cm

# 表格单元格内边距（cm）
CELL_MARGIN_TOP = 0.08
CELL_MARGIN_BOTTOM = 0.08
CELL_MARGIN_LEFT = 0.15
CELL_MARGIN_RIGHT = 0.15

# 默认字体大小
FONT_TABLE_HEADER = 9       # pt
FONT_TABLE_BODY = 9         # pt
FONT_TABLE_SMALL = 8.5      # pt
FONT_LABEL = 9              # pt
FONT_SECTION_LABEL = 10     # pt
FONT_CONCLUSION = 11        # pt

# 差异级别颜色映射
DIFF_COLORS = {
    "ID": GREEN,      # 完全一致
    "MN": GOLD,       # 轻微差异
    "SD": ORANGE,     # 实质差异
    "CD": RED,        # 严重差异
    "UV": MEDIUM_GRAY # 无法核实
}

# 差异级别中文标签
DIFF_LABELS = {
    "ID": "完全一致",
    "MN": "轻微差异",
    "SD": "实质差异",
    "CD": "严重差异",
    "UV": "无法核实"
}

# 状态中文标签
STATUS_LABELS = {
    "current": "现行有效",
    "revised": "部分修订",
    "repealed": "已废止",
    "expired": "已失效",
    "not_yet_effective": "尚未生效"
}

# 状态颜色映射
STATUS_COLORS = {
    "current": GREEN,
    "revised": ORANGE,
    "repealed": RED,
    "expired": RED,
    "not_yet_effective": GOLD
}

# 置信度中文标签
CONFIDENCE_LABELS = {
    "H": "高",
    "M": "中",
    "L": "低"
}


# ============================================================
# 字体检测
# ============================================================

def detect_chinese_font():
    """检测系统中可用的中文字体"""
    import platform
    system = platform.system()

    if system == "Windows":
        heading_font = "微软雅黑"
        body_font = "仿宋"
    elif system == "Darwin":
        heading_font = "PingFang SC"
        body_font = "PingFang SC"
    else:
        heading_font = "Noto Sans CJK SC"
        body_font = "Noto Sans CJK SC"

    return heading_font, body_font


# ============================================================
# DOCX 文档构建器
# ============================================================

class ReportGenerator:
    """法条核验报告 DOCX 生成器"""

    def __init__(self, data: dict, title: str = None):
        self.data = data
        self.title = title or "法条核验报告"
        self.heading_font, self.body_font = detect_chinese_font()

        # 延迟导入
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor, Emu
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            self.Document = Document
            self.Inches = Inches
            self.Pt = Pt
            self.Cm = Cm
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self.WD_TABLE_ALIGNMENT = WD_TABLE_ALIGNMENT
            self.WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL
            self.qn = qn
            self.OxmlElement = OxmlElement
        except ImportError:
            raise ImportError("需要 python-docx 库，请运行：pip install python-docx")

    def generate(self) -> 'Document':
        """生成完整报告"""
        doc = self.Document()

        # 设置默认字体和段落间距
        style = doc.styles['Normal']
        style.font.name = self.body_font
        style.font.size = self.Pt(10.5)
        style.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)
        style.paragraph_format.space_before = self.Pt(2)
        style.paragraph_format.space_after = self.Pt(4)

        # 判断报告模式
        mode = self.data.get('mode', 'single')

        if mode == 'batch':
            self._add_batch_report(doc)
        else:
            self._add_single_report(doc)

        return doc

    # ============================================================
    # 单条核验报告
    # ============================================================

    def _add_single_report(self, doc):
        """添加单条核验报告内容"""

        self._add_title(doc, f"法条核验报告")

        self._add_verification_object(doc)

        self._add_sources(doc)

        self._add_official_text(doc)

        self._add_comparison(doc)

        self._add_validity_status(doc)

        self._add_interpretive_context(doc)

        self._add_jurisdiction(doc)

        self._add_format(doc)

        self._add_confidence(doc)

        self._add_conclusion(doc)

        self._add_disclaimer(doc)

    def _add_verification_object(self, doc):
        """添加核验对象信息"""
        obj = self.data.get('verification_object', {})

        self._add_heading(doc, "核验对象", level=2)

        # 基本信息表 — 两列：标签列 3.5cm + 内容列 12cm = 15.5cm
        row_count = 7
        table = self._create_table(doc, cols=2, rows=row_count, col_widths=[3.5, 12])

        fields = [
            ("法律名称", obj.get('law_name', '—')),
            ("条款", self._format_article_ref(obj)),
            ("颁布机关", obj.get('issuing_authority', '—')),
            ("版本年份", obj.get('version_year', '未指定')),
            ("用户提供内容", obj.get('user_text', '—')),
            ("用户适用意图", obj.get('user_intent', '未提供')),
            ("关键词提取", self._format_keywords(obj.get('keywords', {}))),
        ]

        for i, (label, value) in enumerate(fields):
            row = table.rows[i]
            label_color = DARK_BLUE
            # 用户适用意图行高亮
            if label == "用户适用意图" and value != '未提供':
                label_color = ORANGE
            self._set_cell(row.cells[0], label, bold=True, bg_color=label_color, font_color=WHITE,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(row.cells[1], str(value), v_align=self.WD_ALIGN_VERTICAL.CENTER)

    def _format_keywords(self, keywords: dict) -> str:
        """格式化关键词提取结果"""
        if not keywords:
            return '未提取'
        core = keywords.get('core_terms', [])
        divergence = keywords.get('divergence_points', [])
        parts = []
        if core:
            parts.append(f"核心术语：{', '.join(core)}")
        if divergence:
            parts.append(f"分歧点：{', '.join(divergence)}")
        return '；'.join(parts) if parts else '未提取'

    def _add_sources(self, doc):
        """添加核验来源"""
        sources = self.data.get('sources', [])

        self._add_heading(doc, "核验来源", level=2)

        if not sources:
            self._add_paragraph(doc, "未记录核验来源", italic=True)
            return

        # 四列：来源3.5cm + 网址5.5cm + 状态2.5cm + 备注4cm = 15.5cm
        table = self._create_table(doc, cols=4, rows=len(sources) + 1,
                                   col_widths=[3.5, 5.5, 2.5, 4])

        # 表头
        header = table.rows[0]
        headers = ["来源名称", "网址", "状态", "备注"]
        for i, h in enumerate(headers):
            self._set_cell(header.cells[i], h, bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                          align=self.WD_ALIGN_PARAGRAPH.CENTER, size=self.Pt(FONT_TABLE_HEADER),
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # 数据行
        for i, source in enumerate(sources):
            row = table.rows[i + 1]
            status = source.get('status', 'unknown')
            status_text = "已访问" if status == 'accessed' else "未访问" if status == 'failed' else "部分访问"
            bg = LIGHT_GRAY if i % 2 == 0 else WHITE

            self._set_cell(row.cells[0], source.get('name', '—'), bg_color=bg,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(row.cells[1], source.get('url', '—'), bg_color=bg,
                          size=self.Pt(FONT_TABLE_SMALL),
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(row.cells[2], status_text, bg_color=bg,
                          align=self.WD_ALIGN_PARAGRAPH.CENTER,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(row.cells[3], source.get('note', ''), bg_color=bg,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

    def _add_official_text(self, doc):
        """添加官方原文"""
        official = self.data.get('official_text', {})

        self._add_heading(doc, "官方原文", level=2)

        if not official.get('text'):
            self._add_paragraph(doc, "未能获取官方原文", italic=True, color=RED)
            return

        # 引用块样式 — 带左侧蓝色竖线的引用框
        quote_table = self._create_table(doc, cols=1, rows=1, col_widths=[PAGE_CONTENT_WIDTH])
        cell = quote_table.rows[0].cells[0]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(official['text'])
        run.font.size = self.Pt(FONT_TABLE_BODY)
        run.font.italic = True
        run.font.name = self.body_font
        run.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)
        run.font.color.rgb = self.RGBColor.from_string(DARK_GRAY)
        # 淡蓝背景
        self._set_cell_shading(cell, "EBF5FB")

        # 来源标注
        source_url = official.get('source_url', '')
        if source_url:
            source_para = doc.add_paragraph()
            source_para.paragraph_format.alignment = self.WD_ALIGN_PARAGRAPH.RIGHT
            source_para.paragraph_format.space_before = self.Pt(2)
            source_run = source_para.add_run(f"来源：{source_url}")
            source_run.font.size = self.Pt(FONT_TABLE_SMALL)
            source_run.font.color.rgb = self.RGBColor.from_string(MEDIUM_GRAY)

    def _add_comparison(self, doc):
        """添加内容比对结果"""
        comparison = self.data.get('comparison', {})
        diff_level = comparison.get('diff_level', 'UV')
        differences = comparison.get('differences', [])

        self._add_heading(doc, "内容比对", level=2)

        # 比对结论
        label = DIFF_LABELS.get(diff_level, "未知")
        color = DIFF_COLORS.get(diff_level, MEDIUM_GRAY)

        conclusion_para = doc.add_paragraph()
        conclusion_para.paragraph_format.space_after = self.Pt(6)
        run = conclusion_para.add_run(f"比对结论：{label}")
        run.bold = True
        run.font.size = self.Pt(12)
        run.font.color.rgb = self.RGBColor.from_string(color)

        # 差异详情
        if differences:
            # 五列：位置2cm + 用户版4.5cm + 官方版4.5cm + 差异级别2.5cm + 影响2cm = 15.5cm
            table = self._create_table(doc, cols=5, rows=len(differences) + 1,
                                       col_widths=[2, 4.5, 4.5, 2.5, 2])

            header = table.rows[0]
            headers = ["位置", "用户版", "官方版", "差异级别", "影响说明"]
            for i, h in enumerate(headers):
                self._set_cell(header.cells[i], h, bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER, size=self.Pt(FONT_TABLE_HEADER),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

            for i, diff in enumerate(differences):
                row = table.rows[i + 1]
                diff_type = diff.get('diff_type', 'UV')
                bg = LIGHT_GRAY if i % 2 == 0 else WHITE
                diff_color = DIFF_COLORS.get(diff_type, MEDIUM_GRAY)

                self._set_cell(row.cells[0], diff.get('location', '—'), bg_color=bg,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)
                self._set_cell(row.cells[1], diff.get('user_text', '—'), bg_color=bg,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)
                self._set_cell(row.cells[2], diff.get('official_text', '—'), bg_color=bg,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

                # 差异级别带颜色
                self._set_cell(row.cells[3], DIFF_LABELS.get(diff_type, '—'),
                              bg_color=diff_color, font_color=WHITE,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

                self._set_cell(row.cells[4], diff.get('impact', ''), bg_color=bg,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # 正确法条
        correct = comparison.get('correct_text')
        if correct:
            self._add_spacer(doc, 6)
            correct_para = doc.add_paragraph()
            label_run = correct_para.add_run("正确法条：")
            label_run.bold = True
            label_run.font.color.rgb = self.RGBColor.from_string(GREEN)
            label_run.font.size = self.Pt(FONT_SECTION_LABEL)
            correct_run = correct_para.add_run(correct)
            correct_run.font.size = self.Pt(FONT_SECTION_LABEL)

    def _add_validity_status(self, doc):
        """添加有效性状态"""
        validity = self.data.get('validity', {})

        self._add_heading(doc, "有效性状态", level=2)

        # 两列：标签3.5cm + 内容12cm
        table = self._create_table(doc, cols=2, rows=5, col_widths=[3.5, 12])

        status = validity.get('status', 'unknown')
        status_label = STATUS_LABELS.get(status, '未知')
        status_color = STATUS_COLORS.get(status, MEDIUM_GRAY)

        fields = [
            ("法律状态", status_label),
            ("最新版本", validity.get('latest_version', '—')),
            ("施行日期", validity.get('effective_date', '—')),
            ("该条款状态", validity.get('clause_status', '—')),
            ("替代指引", validity.get('replacement', '无')),
        ]

        for i, (label, value) in enumerate(fields):
            row = table.rows[i]
            self._set_cell(row.cells[0], label, bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

            if i == 0:  # 状态行带颜色
                self._set_cell(row.cells[1], value, font_color=status_color, bold=True,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)
            else:
                self._set_cell(row.cells[1], str(value),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # 替代法律详情
        replacement_detail = validity.get('replacement_detail')
        if replacement_detail:
            self._add_spacer(doc, 6)
            rep_para = doc.add_paragraph()
            label_run = rep_para.add_run("替代法律：")
            label_run.bold = True
            label_run.font.color.rgb = self.RGBColor.from_string(ORANGE)
            rep_para.add_run(replacement_detail)

    def _add_interpretive_context(self, doc):
        """添加司法解释与适用性审查信息"""
        ic = self.data.get('interpretive_context', {})

        self._add_heading(doc, "司法解释与适用性审查", level=2)

        diff_level = ic.get('interpretive_diff_level')
        interpretations = ic.get('judicial_interpretations', [])
        regulations = ic.get('supporting_regulations', [])
        note = ic.get('interpretive_note', '')
        user_intent_analyzed = ic.get('user_intent_analyzed', '')
        keyword_searches = ic.get('keyword_searches_performed', [])

        # 用户适用意图分析
        if user_intent_analyzed:
            intent_para = doc.add_paragraph()
            intent_para.paragraph_format.space_after = self.Pt(4)
            label_run = intent_para.add_run("用户理解分析：")
            label_run.bold = True
            label_run.font.size = self.Pt(FONT_SECTION_LABEL)
            intent_para.add_run(user_intent_analyzed).font.size = self.Pt(FONT_SECTION_LABEL)

        # 搜索执行记录
        if keyword_searches:
            self._add_spacer(doc, 4)
            search_label = doc.add_paragraph()
            search_label.paragraph_format.space_after = self.Pt(4)
            run = search_label.add_run("搜索执行记录：")
            run.bold = True
            run.font.size = self.Pt(FONT_SECTION_LABEL)

            # 四列：轨道3cm + 查询6.5cm + 结果2.5cm + 备注3.5cm = 15.5cm
            search_table = self._create_table(doc, cols=4, rows=len(keyword_searches) + 1,
                                               col_widths=[3, 6.5, 2.5, 3.5])
            header = search_table.rows[0]
            headers = ["搜索轨道", "查询内容", "结果", "备注"]
            for i, h in enumerate(headers):
                self._set_cell(header.cells[i], h, bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER, size=self.Pt(FONT_TABLE_HEADER),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

            for i, search in enumerate(keyword_searches):
                row = search_table.rows[i + 1]
                bg = LIGHT_GRAY if i % 2 == 0 else WHITE
                found = search.get('results_found', False)
                result_text = "找到" if found else "—"
                result_color = GREEN if found else MEDIUM_GRAY

                self._set_cell(row.cells[0], search.get('track', '—'), bg_color=bg,
                              size=self.Pt(FONT_TABLE_BODY),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)
                self._set_cell(row.cells[1], search.get('query', '—'), bg_color=bg,
                              size=self.Pt(FONT_TABLE_BODY),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)
                self._set_cell(row.cells[2], result_text, bg_color=bg, font_color=result_color,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER, size=self.Pt(FONT_TABLE_BODY),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)
                self._set_cell(row.cells[3], search.get('note', ''), bg_color=bg,
                              size=self.Pt(FONT_TABLE_BODY),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # 解释性差异级别
        if diff_level:
            ic_labels = {
                'IC-CD': '严重解释性差异',
                'IC-SD': '实质解释性差异',
                'IC-MN': '轻微解释性差异',
                'IC-Info': '提示性信息',
                'IC-Pending': '发现但未判定（用户未提供适用意图）',
            }
            ic_colors = {
                'IC-CD': RED,
                'IC-SD': ORANGE,
                'IC-MN': GOLD,
                'IC-Info': GREEN,
                'IC-Pending': MEDIUM_GRAY,
            }
            label = ic_labels.get(diff_level, '未知')
            color = ic_colors.get(diff_level, MEDIUM_GRAY)

            self._add_spacer(doc, 4)
            level_para = doc.add_paragraph()
            run = level_para.add_run(f"解释性差异：{label}（{diff_level}）")
            run.bold = True
            run.font.size = self.Pt(12)
            run.font.color.rgb = self.RGBColor.from_string(color)

        # 司法解释列表
        if interpretations:
            self._add_spacer(doc, 4)
            sub_heading = doc.add_paragraph()
            sub_heading.paragraph_format.space_after = self.Pt(4)
            run = sub_heading.add_run("相关司法解释：")
            run.bold = True
            run.font.size = self.Pt(FONT_SECTION_LABEL)

            for interp in interpretations:
                # 两列：标签3cm + 内容12.5cm = 15.5cm
                table = self._create_table(doc, cols=2, rows=6, col_widths=[3, 12.5])
                fields = [
                    ("名称", interp.get('name', '—')),
                    ("文号", interp.get('document_number', '—')),
                    ("施行日期", interp.get('effective_date', '—')),
                    ("相关条款", interp.get('relevant_article', '—')),
                    ("发现轨道", interp.get('found_by_track', '未标注')),
                    ("影响说明", interp.get('impact', '—')),
                ]
                for i, (f_label, value) in enumerate(fields):
                    row = table.rows[i]
                    label_bg = DARK_BLUE
                    # 发现轨道行用绿色高亮
                    if f_label == "发现轨道" and value == '关键词级':
                        label_bg = GREEN
                    self._set_cell(row.cells[0], f_label, bold=True, bg_color=label_bg, font_color=WHITE,
                                  size=self.Pt(FONT_TABLE_BODY),
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)
                    self._set_cell(row.cells[1], str(value), size=self.Pt(FONT_TABLE_BODY),
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)

                # 条文内容
                content = interp.get('content', '')
                if content:
                    content_para = doc.add_paragraph()
                    content_para.paragraph_format.left_indent = self.Cm(0.5)
                    content_para.paragraph_format.space_before = self.Pt(4)
                    content_run = content_para.add_run(f"条文内容：{content}")
                    content_run.font.size = self.Pt(FONT_TABLE_SMALL)
                    content_run.italic = True
        else:
            search_note = doc.add_paragraph()
            search_note.paragraph_format.space_before = self.Pt(4)
            run = search_note.add_run("已执行法条级+关键词级搜索，未发现与该条款直接相关的司法解释")
            run.font.color.rgb = self.RGBColor.from_string(GREEN)
            run.font.size = self.Pt(FONT_TABLE_BODY)

        # 配套规定列表
        if regulations:
            self._add_spacer(doc, 4)
            reg_heading = doc.add_paragraph()
            reg_heading.paragraph_format.space_after = self.Pt(4)
            run = reg_heading.add_run("相关配套规定：")
            run.bold = True
            run.font.size = self.Pt(FONT_SECTION_LABEL)

            for reg in regulations:
                reg_para = doc.add_paragraph()
                reg_para.paragraph_format.left_indent = self.Cm(0.5)
                reg_para.paragraph_format.space_after = self.Pt(2)
                reg_run = reg_para.add_run(f"  {reg.get('name', '—')}：{reg.get('impact', '')}")
                reg_run.font.size = self.Pt(FONT_TABLE_BODY)

        # 解释性差异警告
        if diff_level in ('IC-CD', 'IC-SD'):
            self._add_spacer(doc, 6)
            # 红色醒目警告框
            warn_table = self._create_table(doc, cols=1, rows=1, col_widths=[PAGE_CONTENT_WIDTH])
            cell = warn_table.rows[0].cells[0]
            warn_text = (
                "解释性差异警告：虽然法条文本与官方原文完全一致，"
                "但相关司法解释已改变/明确了该条款的理解方式。"
                "请参考上述司法解释重新审视法律适用。"
            )
            self._set_cell(cell, warn_text, bold=True, font_color=WHITE, bg_color=RED,
                          size=self.Pt(FONT_SECTION_LABEL),
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # IC-Pending提醒
        if diff_level == 'IC-Pending':
            self._add_spacer(doc, 6)
            pending_para = doc.add_paragraph()
            run = pending_para.add_run(
                "发现相关司法解释，但因未提供用户适用意图，无法判定解释性差异级别。"
                "建议说明您基于该条款的法律论点，以便完成判定。"
            )
            run.font.color.rgb = self.RGBColor.from_string(GOLD)
            run.font.size = self.Pt(FONT_SECTION_LABEL)

        # 备注说明
        if note:
            note_para = doc.add_paragraph()
            note_para.paragraph_format.space_before = self.Pt(4)
            note_run = note_para.add_run(f"备注：{note}")
            note_run.font.italic = True
            note_run.font.color.rgb = self.RGBColor.from_string(MEDIUM_GRAY)
            note_run.font.size = self.Pt(FONT_TABLE_SMALL)

    def _add_jurisdiction(self, doc):
        """添加管辖区适配检查"""
        jurisdiction = self.data.get('jurisdiction', {})

        self._add_heading(doc, "管辖区适配", level=2)

        inferred = jurisdiction.get('inferred', '未指定')
        basis = jurisdiction.get('inference_basis', '')
        issues = jurisdiction.get('issues', [])

        info_table = self._create_table(doc, cols=2, rows=2 if basis else 1, col_widths=[3.5, 12])

        self._set_cell(info_table.rows[0].cells[0], "适用管辖区", bold=True, bg_color=DARK_BLUE,
                      font_color=WHITE, v_align=self.WD_ALIGN_VERTICAL.CENTER)
        self._set_cell(info_table.rows[0].cells[1], inferred,
                      v_align=self.WD_ALIGN_VERTICAL.CENTER)

        if basis:
            self._set_cell(info_table.rows[1].cells[0], "推断依据", bold=True, bg_color=DARK_BLUE,
                          font_color=WHITE, v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(info_table.rows[1].cells[1], basis,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

        if issues:
            for issue in issues:
                issue_type = issue.get('type', '')
                description = issue.get('description', '')
                severity = issue.get('severity', '🟡')
                issue_para = doc.add_paragraph()
                issue_para.paragraph_format.space_before = self.Pt(4)
                run = issue_para.add_run(f"{severity} {issue_type}：{description}")
                run.font.size = self.Pt(FONT_TABLE_BODY)
                if severity == '🔴':
                    run.font.color.rgb = self.RGBColor.from_string(RED)
                    run.bold = True
                elif severity == '🟠':
                    run.font.color.rgb = self.RGBColor.from_string(ORANGE)
                else:
                    run.font.color.rgb = self.RGBColor.from_string(GOLD)
        else:
            ok_para = doc.add_paragraph()
            ok_para.paragraph_format.space_before = self.Pt(4)
            run = ok_para.add_run("未发现管辖区适配问题")
            run.font.color.rgb = self.RGBColor.from_string(GREEN)
            run.font.size = self.Pt(FONT_TABLE_BODY)

    def _add_format(self, doc):
        """添加格式规范化建议"""
        fmt = self.data.get('format', {})

        self._add_heading(doc, "格式规范化", level=2)

        original = fmt.get('original', '—')
        normalized = fmt.get('normalized', '—')
        issues = fmt.get('issues', [])

        fmt_table = self._create_table(doc, cols=2, rows=2, col_widths=[3.5, 12])

        self._set_cell(fmt_table.rows[0].cells[0], "原引用格式", bold=True, bg_color=DARK_BLUE,
                      font_color=WHITE, v_align=self.WD_ALIGN_VERTICAL.CENTER)
        self._set_cell(fmt_table.rows[0].cells[1], original,
                      v_align=self.WD_ALIGN_VERTICAL.CENTER)

        self._set_cell(fmt_table.rows[1].cells[0], "规范格式", bold=True, bg_color=DARK_BLUE,
                      font_color=WHITE, v_align=self.WD_ALIGN_VERTICAL.CENTER)
        self._set_cell(fmt_table.rows[1].cells[1], normalized,
                      v_align=self.WD_ALIGN_VERTICAL.CENTER)

        if issues:
            self._add_spacer(doc, 4)
            issues_para = doc.add_paragraph()
            issues_para.paragraph_format.space_after = self.Pt(3)
            label_run = issues_para.add_run("格式问题：")
            label_run.bold = True
            label_run.font.size = self.Pt(FONT_SECTION_LABEL)
            for issue in issues:
                item_para = doc.add_paragraph()
                item_para.paragraph_format.left_indent = self.Cm(0.5)
                item_para.paragraph_format.space_after = self.Pt(2)
                run = item_para.add_run(f"🟡 {issue}")
                run.font.size = self.Pt(FONT_TABLE_BODY)
                run.font.color.rgb = self.RGBColor.from_string(GOLD)
        else:
            ok_para = doc.add_paragraph()
            ok_para.paragraph_format.space_before = self.Pt(4)
            run = ok_para.add_run("引用格式符合规范")
            run.font.color.rgb = self.RGBColor.from_string(GREEN)
            run.font.size = self.Pt(FONT_TABLE_BODY)

    def _add_confidence(self, doc):
        """添加置信度评估"""
        confidence = self.data.get('confidence', {})
        level = confidence.get('confidence_level', 'L')
        factors = confidence.get('confidence_factors', {})
        note = confidence.get('confidence_note', '')

        self._add_heading(doc, "置信度评估", level=2)

        level_label = CONFIDENCE_LABELS.get(level, '未知')
        level_color = {
            'H': GREEN, 'M': ORANGE, 'L': RED
        }.get(level, MEDIUM_GRAY)

        para = doc.add_paragraph()
        para.paragraph_format.space_after = self.Pt(6)
        run = para.add_run(f"置信度等级：{level_label}（{level}）")
        run.bold = True
        run.font.size = self.Pt(12)
        run.font.color.rgb = self.RGBColor.from_string(level_color)

        # 因子详情
        if factors:
            factor_labels = {
                'official_sources_count': '官方来源数量',
                'sources_agree': '来源内容一致',
                'current_status_confirmed': '现行状态已确认',
                'latest_version_confirmed': '最新版本已确认',
                'word_by_word_comparison_done': '逐字比对已完成',
                'proviso_checked': '但书/除外条款已检查',
                'judicial_interpretation_searched': '司法解释法条级搜索',
                'keyword_level_search_performed': '司法解释关键词级搜索',
                'supporting_regulations_searched': '配套规定已检索',
                'user_intent_confirmed': '用户适用意图已确认',
            }

            # 两列：因子6.5cm + 结果9cm = 15.5cm
            table = self._create_table(doc, cols=2, rows=len(factors) + 1,
                                       col_widths=[6.5, 9])

            header = table.rows[0]
            self._set_cell(header.cells[0], "评估因子", bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                          size=self.Pt(FONT_TABLE_HEADER),
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(header.cells[1], "结果", bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                          align=self.WD_ALIGN_PARAGRAPH.CENTER, size=self.Pt(FONT_TABLE_HEADER),
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

            for i, (key, value) in enumerate(factors.items()):
                row = table.rows[i + 1]
                bg = LIGHT_GRAY if i % 2 == 0 else WHITE
                f_label = factor_labels.get(key, key)

                self._set_cell(row.cells[0], f_label, bg_color=bg,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

                if isinstance(value, bool):
                    result = "是" if value else "否"
                else:
                    result = str(value)
                self._set_cell(row.cells[1], result, bg_color=bg,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

        if note:
            note_para = doc.add_paragraph()
            note_para.paragraph_format.space_before = self.Pt(4)
            note_run = note_para.add_run(f"备注：{note}")
            note_run.font.italic = True
            note_run.font.color.rgb = self.RGBColor.from_string(MEDIUM_GRAY)

        # 低置信度警告
        if level == 'L':
            self._add_spacer(doc, 6)
            warn_para = doc.add_paragraph()
            warn_run = warn_para.add_run("本核验置信度为低，建议通过官方渠道二次确认")
            warn_run.bold = True
            warn_run.font.color.rgb = self.RGBColor.from_string(RED)
        elif level == 'M':
            self._add_spacer(doc, 6)
            warn_para = doc.add_paragraph()
            warn_run = warn_para.add_run("本核验置信度为中等，建议通过官方渠道二次确认")
            warn_run.font.color.rgb = self.RGBColor.from_string(ORANGE)

    def _add_conclusion(self, doc):
        """添加核验结论"""
        conclusion = self.data.get('conclusion', '')

        self._add_heading(doc, "核验结论", level=2)

        para = doc.add_paragraph()
        run = para.add_run(conclusion)
        run.font.size = self.Pt(FONT_CONCLUSION)

    # ============================================================
    # 批量核验报告
    # ============================================================

    def _add_batch_report(self, doc):
        """添加批量核验报告内容"""

        self._add_title(doc, "法条核验报告")

        # 文件概况
        overview = self.data.get('overview', {})
        self._add_heading(doc, "文件概况", level=2)

        severe = overview.get('severe_issues', 0)
        important = overview.get('important_issues', 0)
        minor = overview.get('minor_issues', 0)

        table = self._create_table(doc, cols=2, rows=5, col_widths=[4, 11.5])
        fields = [
            ("文件名称", overview.get('file_name', '—')),
            ("法律引用数量", str(overview.get('total_citations', 0))),
            ("核验完成", f"{overview.get('passed', 0)} 条通过 / "
                        f"{overview.get('issues', 0)} 条存在问题 / "
                        f"{overview.get('unverified', 0)} 条无法核实"),
            ("问题分布", f"🔴 严重 {severe} 条 / 🟠 重要 {important} 条 / 🟡 轻微 {minor} 条"),
            ("核验日期", datetime.now().strftime('%Y-%m-%d')),
        ]
        for i, (label, value) in enumerate(fields):
            row = table.rows[i]
            self._set_cell(row.cells[0], label, bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)
            self._set_cell(row.cells[1], str(value),
                          v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # 核验结果概览 — 使用卡片式布局，每个法条一张表
        items = self.data.get('items', [])
        self._add_heading(doc, "核验结果概览", level=2)

        if items:
            ic_labels_map = {
                'IC-CD': '严重',
                'IC-SD': '实质',
                'IC-MN': '轻微',
                'IC-Info': '提示',
                'IC-Pending': '待判定',
            }
            ic_colors_map = {
                'IC-CD': RED,
                'IC-SD': ORANGE,
                'IC-MN': GOLD,
                'IC-Info': GREEN,
                'IC-Pending': MEDIUM_GRAY,
            }

            overview_table = self._create_table(
                doc, cols=7, rows=len(items) + 1,
                col_widths=[0.8, 4.5, 2.5, 2, 2, 1.8, 1.9]
            )

            header = overview_table.rows[0]
            headers = ["序号", "法律+条款", "核验/有效性", "解释性差异", "管辖区", "格式", "说明"]
            for i, h in enumerate(headers):
                self._set_cell(header.cells[i], h, bold=True, bg_color=DARK_BLUE, font_color=WHITE,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER, size=self.Pt(FONT_TABLE_HEADER),
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

            for i, item in enumerate(items):
                row = overview_table.rows[i + 1]
                bg = LIGHT_GRAY if i % 2 == 0 else WHITE

                diff_level = item.get('diff_level', 'UV')
                validity = item.get('validity', 'unknown')
                ic_level = item.get('interpretive_diff_level')

                diff_label = DIFF_LABELS.get(diff_level, '—')
                diff_color = DIFF_COLORS.get(diff_level, MEDIUM_GRAY)
                valid_label = STATUS_LABELS.get(validity, '—')
                valid_color = STATUS_COLORS.get(validity, MEDIUM_GRAY)

                # 合并法律名称+条款
                law_article = f"{item.get('law_name', '—')}\n{item.get('article_ref', '')}"

                # 合并内容核验+有效性 — 两行显示
                check_validity = f"{diff_label} / {valid_label}"

                self._set_cell(row.cells[0], str(i + 1), bg_color=bg,
                              align=self.WD_ALIGN_PARAGRAPH.CENTER,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

                # 法律+条款列 — 左对齐，法律名称普通+条款小字
                cell1 = row.cells[1]
                cell1.text = ''
                p1 = cell1.paragraphs[0]
                r1 = p1.add_run(item.get('law_name', '—'))
                r1.font.size = self.Pt(FONT_TABLE_BODY)
                r1.font.name = self.body_font
                r1.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)
                p2 = cell1.add_paragraph()
                r2 = p2.add_run(item.get('article_ref', '—'))
                r2.font.size = self.Pt(FONT_TABLE_SMALL)
                r2.font.color.rgb = self.RGBColor.from_string(MEDIUM_GRAY)
                r2.font.name = self.body_font
                r2.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)
                if bg == LIGHT_GRAY:
                    self._set_cell_shading(cell1, LIGHT_GRAY)

                # 内容核验+有效性 — 用两行色块
                cell2 = row.cells[2]
                cell2.text = ''
                # 第一行：内容核验
                p_diff = cell2.paragraphs[0]
                p_diff.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
                r_diff = p_diff.add_run(diff_label)
                r_diff.bold = True
                r_diff.font.size = self.Pt(FONT_TABLE_SMALL)
                r_diff.font.color.rgb = self.RGBColor.from_string(diff_color)
                r_diff.font.name = self.body_font
                r_diff.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)
                # 第二行：有效性
                p_val = cell2.add_paragraph()
                p_val.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
                r_val = p_val.add_run(valid_label)
                r_val.font.size = self.Pt(FONT_TABLE_SMALL)
                r_val.font.color.rgb = self.RGBColor.from_string(valid_color)
                r_val.font.name = self.body_font
                r_val.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)
                if bg == LIGHT_GRAY:
                    self._set_cell_shading(cell2, LIGHT_GRAY)
                self._set_cell_vertical_align(cell2, self.WD_ALIGN_VERTICAL.CENTER)

                if ic_level:
                    ic_label = ic_labels_map.get(ic_level, '—')
                    ic_color = ic_colors_map.get(ic_level, MEDIUM_GRAY)
                    self._set_cell(row.cells[3], ic_label, bg_color=ic_color, font_color=WHITE,
                                  align=self.WD_ALIGN_PARAGRAPH.CENTER,
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)
                else:
                    self._set_cell(row.cells[3], '—', bg_color=bg,
                                  align=self.WD_ALIGN_PARAGRAPH.CENTER,
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)

                jur_issue = item.get('jurisdiction_issue')
                if jur_issue:
                    self._set_cell(row.cells[4], '⚠', bg_color=ORANGE, font_color=WHITE,
                                  align=self.WD_ALIGN_PARAGRAPH.CENTER,
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)
                else:
                    self._set_cell(row.cells[4], '✅', bg_color=bg,
                                  align=self.WD_ALIGN_PARAGRAPH.CENTER,
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)

                fmt_issue = item.get('format_issue')
                if fmt_issue:
                    self._set_cell(row.cells[5], '🟡', bg_color=GOLD, font_color=WHITE,
                                  align=self.WD_ALIGN_PARAGRAPH.CENTER,
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)
                else:
                    self._set_cell(row.cells[5], '✅', bg_color=bg,
                                  align=self.WD_ALIGN_PARAGRAPH.CENTER,
                                  v_align=self.WD_ALIGN_VERTICAL.CENTER)

                self._set_cell(row.cells[6], item.get('note', ''), bg_color=bg,
                              v_align=self.WD_ALIGN_VERTICAL.CENTER)

        # 详细核验信息 — 每个条目用卡片式双列表
        detail_items = [item for item in items
                       if item.get('diff_level', 'ID') != 'ID'
                       or item.get('validity') != 'current'
                       or item.get('interpretive_diff_level')]
        if detail_items:
            self._add_heading(doc, "详细核验信息", level=2)
            for idx, item in enumerate(detail_items):
                # 条目标题
                item_title = f"{item.get('law_name', '—')} {item.get('article_ref', '')}"
                self._add_heading(doc, item_title, level=3)

                # 卡片式信息表 — 两列
                card_fields = []

                if item.get('differences'):
                    for diff in item['differences']:
                        card_fields.append(("文本差异",
                            f"用户版：{diff.get('user_text', '')}\n官方版：{diff.get('official_text', '')}"))

                if item.get('validity_note'):
                    card_fields.append(("有效性", item['validity_note']))

                if item.get('judicial_interpretations'):
                    for interp in item['judicial_interpretations']:
                        card_fields.append(("司法解释",
                            f"{interp.get('name', '—')} — {interp.get('impact', '')}"))

                if item.get('interpretive_diff_level'):
                    ic_lev = item['interpretive_diff_level']
                    ic_labels_full = {
                        'IC-CD': '严重解释性差异',
                        'IC-SD': '实质解释性差异',
                        'IC-MN': '轻微解释性差异',
                        'IC-Info': '提示性信息',
                        'IC-Pending': '待判定',
                    }
                    card_fields.append(("解释性差异", ic_labels_full.get(ic_lev, ic_lev)))

                if item.get('user_intent'):
                    card_fields.append(("用户适用意图", item['user_intent']))

                if item.get('format_issue'):
                    card_fields.append(("格式问题", item['format_issue']))

                if item.get('jurisdiction_issue'):
                    card_fields.append(("管辖区问题", item['jurisdiction_issue']))

                if card_fields:
                    card_table = self._create_table(doc, cols=2, rows=len(card_fields),
                                                     col_widths=[3.5, 12])
                    for j, (c_label, c_value) in enumerate(card_fields):
                        c_row = card_table.rows[j]
                        self._set_cell(c_row.cells[0], c_label, bold=True, bg_color=DARK_BLUE,
                                      font_color=WHITE, size=self.Pt(FONT_TABLE_BODY),
                                      v_align=self.WD_ALIGN_VERTICAL.CENTER)
                        self._set_cell(c_row.cells[1], c_value, size=self.Pt(FONT_TABLE_BODY),
                                      v_align=self.WD_ALIGN_VERTICAL.CENTER)

                # 卡片之间加间距
                if idx < len(detail_items) - 1:
                    self._add_spacer(doc, 8)

        # 总结建议
        summary = self.data.get('summary', {})
        self._add_heading(doc, "总结建议", level=2)

        categories = {
            'continue': ('可继续使用', GREEN),
            'needs_correction': ('需修正内容', RED),
            'needs_replacement': ('需更换法律', ORANGE),
            'needs_reinterpretation': ('需重新理解', ORANGE),
            'needs_jurisdiction_fix': ('需调整管辖区引用', ORANGE),
        }

        has_summary = False
        for key, (s_label, color) in categories.items():
            items_list = summary.get(key, [])
            if items_list:
                has_summary = True
                para = doc.add_paragraph()
                para.paragraph_format.space_after = self.Pt(3)
                label_run = para.add_run(f"{s_label}：")
                label_run.bold = True
                label_run.font.color.rgb = self.RGBColor.from_string(color)
                label_run.font.size = self.Pt(FONT_SECTION_LABEL)
                content_run = para.add_run('、'.join(str(x) for x in items_list))
                content_run.font.size = self.Pt(FONT_SECTION_LABEL)

        if not has_summary:
            self._add_paragraph(doc, "所有法条引用均无重大问题", italic=True, color=GREEN)

        # 格式问题汇总
        format_fixes = summary.get('needs_format_fix', [])
        if format_fixes:
            self._add_spacer(doc, 4)
            format_para = doc.add_paragraph()
            format_label = format_para.add_run("格式修正：")
            format_label.bold = True
            format_label.font.color.rgb = self.RGBColor.from_string(GOLD)
            format_label.font.size = self.Pt(FONT_SECTION_LABEL)
            format_content = format_para.add_run('；'.join(str(x) for x in format_fixes))
            format_content.font.size = self.Pt(FONT_SECTION_LABEL)

        # 免责声明
        self._add_disclaimer(doc)

    # ============================================================
    # 通用组件
    # ============================================================

    def _add_title(self, doc, text: str):
        """添加报告标题"""
        para = doc.add_paragraph()
        para.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = self.Pt(2)
        run = para.add_run(text)
        run.bold = True
        run.font.size = self.Pt(18)
        run.font.color.rgb = self.RGBColor.from_string(DARK_BLUE)
        run.font.name = self.heading_font
        run.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.heading_font)

        # 核验日期
        date_para = doc.add_paragraph()
        date_para.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_after = self.Pt(12)
        date_run = date_para.add_run(f"核验日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_run.font.size = self.Pt(10)
        date_run.font.color.rgb = self.RGBColor.from_string(MEDIUM_GRAY)

    def _add_heading(self, doc, text: str, level: int = 2):
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = self.RGBColor.from_string(DARK_BLUE)
            run.font.name = self.heading_font
            run.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.heading_font)
        # 标题段前间距
        heading.paragraph_format.space_before = self.Pt(12) if level == 2 else self.Pt(8)
        heading.paragraph_format.space_after = self.Pt(6)

    def _add_paragraph(self, doc, text: str, bold: bool = False, italic: bool = False,
                       color: str = None, size=None):
        """添加段落"""
        para = doc.add_paragraph()
        if text:
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = self.RGBColor.from_string(color)
            if size:
                run.font.size = size
        return para

    def _add_spacer(self, doc, pts: int = 6):
        """添加间距（用零宽空段实现）"""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = self.Pt(0)
        spacer.paragraph_format.space_after = self.Pt(0)
        spacer.paragraph_format.line_spacing = self.Pt(pts)

    def _add_disclaimer(self, doc):
        """添加免责声明"""
        self._add_spacer(doc, 12)

        # 分割线
        para = doc.add_paragraph()
        pBdr = self.OxmlElement('w:pBdr')
        bottom = self.OxmlElement('w:bottom')
        bottom.set(self.qn('w:val'), 'single')
        bottom.set(self.qn('w:sz'), '6')
        bottom.set(self.qn('w:space'), '1')
        bottom.set(self.qn('w:color'), GOLD)
        pBdr.append(bottom)
        para.paragraph_format.element.get_or_add_pPr().append(pBdr)

        disclaimer = doc.add_paragraph()
        disclaimer.paragraph_format.space_before = self.Pt(4)
        run = disclaimer.add_run("免责声明：本核验基于公开信息，仅供参考，不构成法律意见。具体法律问题请咨询专业律师。")
        run.font.size = self.Pt(8)
        run.font.color.rgb = self.RGBColor.from_string(MEDIUM_GRAY)
        run.italic = True

    def _create_table(self, doc, cols: int, rows: int, col_widths: list = None):
        """创建表格并设置列宽、内边距、自动调整"""
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = self.WD_TABLE_ALIGNMENT.CENTER

        # 关闭自动调整，确保列宽生效
        table.autofit = False

        # 设置列宽
        if col_widths and len(col_widths) == cols:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = self.Cm(width)

        # 设置所有单元格的内边距和垂直对齐
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_margins(cell)
                self._set_cell_vertical_align(cell, self.WD_ALIGN_VERTICAL.CENTER)

        return table

    def _set_cell_margins(self, cell):
        """设置单元格内边距"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = self.OxmlElement('w:tcMar')

        for side, value in [
            ('top', CELL_MARGIN_TOP),
            ('bottom', CELL_MARGIN_BOTTOM),
            ('left', CELL_MARGIN_LEFT),
            ('right', CELL_MARGIN_RIGHT),
        ]:
            elem = self.OxmlElement(f'w:{side}')
            elem.set(self.qn('w:w'), str(int(value * 567)))  # cm → twips (1cm = 567 twips)
            elem.set(self.qn('w:type'), 'dxa')
            tcMar.append(elem)

        # 移除已有 tcMar 再添加
        existing = tcPr.find(self.qn('w:tcMar'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcMar)

    def _set_cell_vertical_align(self, cell, align):
        """设置单元格垂直对齐 — align 可以是 WD_ALIGN_VERTICAL 枚举或字符串"""
        # 将枚举转为字符串值
        if hasattr(align, 'value'):
            align_str = str(align)  # WD_ALIGN_VERTICAL.CENTER → 'CENTER (1)'
            # python-docx 枚举的 XML 映射
            align_map = {
                'TOP (0)': 'top',
                'CENTER (1)': 'center',
                'BOTTOM (2)': 'bottom',
                'BOTH (3)': 'both',
            }
            align_val = align_map.get(align_str, 'center')
        elif isinstance(align, str):
            align_val = align
        else:
            align_val = 'center'

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = self.OxmlElement('w:vAlign')
        vAlign.set(self.qn('w:val'), align_val)
        # 移除已有 vAlign
        existing = tcPr.find(self.qn('w:vAlign'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(vAlign)

    def _set_cell_shading(self, cell, color: str):
        """设置单元格背景色"""
        shading = self.OxmlElement('w:shd')
        shading.set(self.qn('w:fill'), color)
        shading.set(self.qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell(self, cell, text: str, bold: bool = False, bg_color: str = None,
                  font_color: str = None, align=None, size=None, v_align=None):
        """设置单元格内容和样式"""
        cell.text = ''
        para = cell.paragraphs[0]

        # 段内间距
        para.paragraph_format.space_before = self.Pt(1)
        para.paragraph_format.space_after = self.Pt(1)
        para.paragraph_format.line_spacing = self.Pt(14)  # 固定行高14pt避免行高跳跃

        if align:
            para.alignment = align

        run = para.add_run(str(text))
        run.bold = bold
        run.font.name = self.body_font
        run.element.rPr.rFonts.set(self.qn('w:eastAsia'), self.body_font)

        if font_color:
            run.font.color.rgb = self.RGBColor.from_string(font_color)
        if size:
            run.font.size = size
        else:
            run.font.size = self.Pt(FONT_TABLE_BODY)

        if bg_color:
            self._set_cell_shading(cell, bg_color)

        if v_align:
            self._set_cell_vertical_align(cell, v_align)

    def _format_article_ref(self, obj: dict) -> str:
        """格式化条款引用"""
        parts = []
        if obj.get('article'):
            parts.append(f"第{obj['article']}条")
        if obj.get('paragraph'):
            parts.append(f"第{obj['paragraph']}款")
        if obj.get('item'):
            parts.append(f"第{obj['item']}项")
        return ''.join(parts) if parts else '—'


# ============================================================
# 主函数
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='法条核验报告 DOCX 生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python generate_report.py result.json
  python generate_report.py result.json --output report.docx
  python generate_report.py result.json --title "《劳动合同法》第47条核验报告"
        """,
    )
    parser.add_argument('json_file', help='核验结果 JSON 文件路径')
    parser.add_argument('-o', '--output', help='输出 DOCX 文件路径（默认与 JSON 同名 .docx）')
    parser.add_argument('-t', '--title', help='报告标题', default='法条核验报告')

    args = parser.parse_args()

    # 读取 JSON
    json_path = Path(args.json_file)
    if not json_path.exists():
        print(f"错误：文件不存在：{args.json_file}", file=sys.stderr)
        sys.exit(1)

    try:
        data = json.loads(json_path.read_text(encoding='utf-8'))
    except json.JSONDecodeError as e:
        print(f"错误：JSON 解析失败：{e}", file=sys.stderr)
        sys.exit(1)

    # 生成报告
    try:
        generator = ReportGenerator(data, title=args.title)
        doc = generator.generate()
    except ImportError as e:
        print(f"错误：{e}", file=sys.stderr)
        sys.exit(1)

    # 输出路径
    output_path = args.output or str(json_path.with_suffix('.docx'))

    # 保存
    doc.save(output_path)
    print(f"报告已保存到：{output_path}", file=sys.stderr)


if __name__ == '__main__':
    main()
