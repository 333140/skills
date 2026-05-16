#!/usr/bin/env python3
"""
法律引用批量提取器 (parse_citations.py)

从文件中提取法律引用，输出标准化 JSON 格式。

用法：
    python parse_citations.py <文件路径> [--output <输出路径>]

支持的文件格式：.docx, .pdf, .txt, .md
"""

import argparse
import json
import re
import sys
from pathlib import Path
from collections import OrderedDict


# ============================================================
# 法律引用正则模式
# ============================================================

CITATION_PATTERNS = [
    (r'(?:依据|根据|按照|依照|遵照)《([^》]+)》第(\d+)条(?:第(\d+)款)?(?:第(\d+)项)?',
     '依据模式'),

    (r'《([^》]+)》第(\d+)条(?:第(\d+)款)?(?:第(\d+)项)?',
     '标准引用'),

    (r'《([^》]+)》第(\d+)条之([一二三四五六七八九十\d]+)(?:第(\d+)款)?(?:第(\d+)项)?',
     '修正案条款引用'),

    (r'《([^》]+)》第([一二三四五六七八九十\d]+)编(?:第([一二三四五六七八九十\d]+)章)?(?:第([一二三四五六七八九十\d]+)节)?',
     '编章节引用'),

    (r'《([^》]+)》第(\d+)款',
     '款引用'),

    (r'《([^》]+)》第(\d+)项',
     '项引用'),

    (r'《([^》]+)》(?:的)?(?:相关)?规定',
     '泛引用'),

    (r'《([^》]+)》第(\d+)条(?:第(\d+)款)?至第(\d+)条(?:第(\d+)款)?',
     '范围引用'),

    (r'([\u4e00-\u9fff]{1,8}(?:法|典|条例|规定|办法|决定|意见|通知|规则|标准|细则|解释))第(\d+)条之([一二三四五六七八九十\d]+)(?:第(\d+)款)?(?:第(\d+)项)?',
     '无书名号修正案引用'),

    (r'([\u4e00-\u9fff]{1,8}(?:法|典|条例|规定|办法|决定|意见|通知|规则|标准|细则|解释))第(\d+)条(?:第(\d+)款)?(?:第(\d+)项)?',
     '无书名号引用'),

    (r'《([^》]+)》法释〔(\d{4})〕(\d+)号',
     '司法解释号引用'),
]


# ============================================================
# 文件读取函数
# ============================================================

def read_txt(file_path: Path) -> str:
    """读取纯文本文件"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
    for enc in encodings:
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法解码文件：{file_path}")


def read_docx(file_path: Path) -> str:
    """读取 DOCX 文件文本内容"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要 python-docx 库，请运行：pip install python-docx")

    doc = Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    # 提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return '\n'.join(paragraphs)


def read_pdf(file_path: Path) -> str:
    """读取 PDF 文件文本内容"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("需要 pdfplumber 库，请运行：pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return '\n'.join(text_parts)


def read_markdown(file_path: Path) -> str:
    """读取 Markdown 文件（等同于文本文件）"""
    return read_txt(file_path)


def read_file(file_path: str) -> str:
    """根据文件扩展名选择读取方式"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{file_path}")

    suffix = path.suffix.lower()
    readers = {
        '.txt': read_txt,
        '.md': read_markdown,
        '.docx': read_docx,
        '.pdf': read_pdf,
    }

    reader = readers.get(suffix)
    if reader is None:
        raise ValueError(f"不支持的文件格式：{suffix}，支持：{', '.join(readers.keys())}")

    return reader(path)


# ============================================================
# 引用提取函数
# ============================================================

def extract_citations(text: str) -> list[dict]:
    """从文本中提取所有法律引用"""
    citations = []

    for pattern, pattern_name in CITATION_PATTERNS:
        for match in re.finditer(pattern, text):
            groups = match.groups()
            citation = {
                'raw_text': match.group(0),
                'pattern_type': pattern_name,
                'position': match.start(),
            }

            # 根据模式类型解析字段
            if pattern_name == '依据模式' or pattern_name == '标准引用':
                citation['law_name'] = groups[0]
                citation['article'] = int(groups[1]) if groups[1] else None
                citation['paragraph'] = int(groups[2]) if groups[2] else None
                citation['item'] = int(groups[3]) if groups[3] else None

            elif pattern_name == '修正案条款引用':
                citation['law_name'] = groups[0]
                citation['article'] = int(groups[1]) if groups[1] else None
                citation['article_suffix'] = groups[2] if groups[2] else None
                citation['paragraph'] = int(groups[3]) if groups[3] else None
                citation['item'] = int(groups[4]) if groups[4] else None

            elif pattern_name == '编章节引用':
                citation['law_name'] = groups[0]
                citation['bian'] = groups[1] if groups[1] else None
                citation['zhang'] = groups[2] if groups[2] else None
                citation['jie'] = groups[3] if groups[3] else None

            elif pattern_name == '款引用':
                citation['law_name'] = groups[0]
                citation['paragraph'] = int(groups[1]) if groups[1] else None

            elif pattern_name == '项引用':
                citation['law_name'] = groups[0]
                citation['item'] = int(groups[1]) if groups[1] else None

            elif pattern_name == '泛引用':
                citation['law_name'] = groups[0]
                citation['reference_type'] = 'general'

            elif pattern_name == '范围引用':
                citation['law_name'] = groups[0]
                citation['article_start'] = int(groups[1]) if groups[1] else None
                citation['paragraph_start'] = int(groups[2]) if groups[2] else None
                citation['article_end'] = int(groups[3]) if groups[3] else None
                citation['paragraph_end'] = int(groups[4]) if groups[4] else None

            elif pattern_name == '无书名号引用':
                citation['law_name'] = groups[0]
                citation['article'] = int(groups[1]) if groups[1] else None
                citation['paragraph'] = int(groups[2]) if groups[2] else None
                citation['item'] = int(groups[3]) if groups[3] else None

            elif pattern_name == '无书名号修正案引用':
                citation['law_name'] = groups[0]
                citation['article'] = int(groups[1]) if groups[1] else None
                citation['article_suffix'] = groups[2] if groups[2] else None
                citation['paragraph'] = int(groups[3]) if groups[3] else None
                citation['item'] = int(groups[4]) if groups[4] else None

            elif pattern_name == '司法解释号引用':
                citation['law_name'] = groups[0]
                citation['judicial_year'] = int(groups[1]) if groups[1] else None
                citation['judicial_number'] = int(groups[2]) if groups[2] else None

            citations.append(citation)

    # 按位置排序
    citations.sort(key=lambda x: x['position'])

    return citations


def normalize_law_name(name: str) -> str:
    """标准化法律名称（添加书名号、补全前缀等）"""
    # 如果没有书名号，添加
    if not name.startswith('《'):
        name = f'《{name}》'

    # 常见简称→全称映射
    common_mappings = {
        '《劳动法》': '《中华人民共和国劳动法》',
        '《劳动合同法》': '《中华人民共和国劳动合同法》',
        '《民法典》': '《中华人民共和国民法典》',
        '《公司法》': '《中华人民共和国公司法》',
        '《刑法》': '《中华人民共和国刑法》',
        '《行政法》': '《中华人民共和国行政强制法》',
        '《行政诉讼法》': '《中华人民共和国行政诉讼法》',
        '《民事诉讼法》': '《中华人民共和国民事诉讼法》',
        '《仲裁法》': '《中华人民共和国仲裁法》',
        '《商标法》': '《中华人民共和国商标法》',
        '《专利法》': '《中华人民共和国专利法》',
        '《著作权法》': '《中华人民共和国著作权法》',
        '《反垄断法》': '《中华人民共和国反垄断法》',
        '《证券法》': '《中华人民共和国证券法》',
        '《税法》': None,
        '《宪法》': '《中华人民共和国宪法》',
        '《行政处罚法》': '《中华人民共和国行政处罚法》',
        '《行政复议法》': '《中华人民共和国行政复议法》',
        '《安全生产法》': '《中华人民共和国安全生产法》',
        '《数据安全法》': '《中华人民共和国数据安全法》',
        '《个保法》': '《中华人民共和国个人信息保护法》',
        '《个人信息保护法》': '《中华人民共和国个人信息保护法》',
        '民法典': '《中华人民共和国民法典》',
        '刑法': '《中华人民共和国刑法》',
    }

    return common_mappings.get(name, name)


def group_by_law(citations: list[dict]) -> dict:
    """按法律名称分组去重"""
    groups = OrderedDict()

    for citation in citations:
        law_name = citation.get('law_name', '未知法律')
        normalized = normalize_law_name(law_name)

        if normalized not in groups:
            groups[normalized] = {
                'law_name': normalized,
                'original_name': law_name,
                'citations': [],
                'articles_mentioned': set(),
            }

        # 提取条款编号用于去重
        article = citation.get('article')
        if article:
            groups[normalized]['articles_mentioned'].add(article)

        # 去重：同一法律同一条款同一位置不重复添加
        article_key = (
            article,
            citation.get('article_suffix'),
            citation.get('paragraph'),
            citation.get('item'),
            citation.get('bian'),
            citation.get('zhang'),
            citation.get('jie'),
        )
        existing_keys = [
            (
                c.get('article'),
                c.get('article_suffix'),
                c.get('paragraph'),
                c.get('item'),
                c.get('bian'),
                c.get('zhang'),
                c.get('jie'),
            )
            for c in groups[normalized]['citations']
        ]
        if article_key not in existing_keys:
            groups[normalized]['citations'].append(citation)

    # 将 set 转为 sorted list 以便 JSON 序列化
    for group in groups.values():
        group['articles_mentioned'] = sorted(group['articles_mentioned'])
        group['citation_count'] = len(group['citations'])

    return groups


def generate_output(citations: list[dict], groups: dict, source_file: str) -> dict:
    """生成最终输出 JSON"""
    # 统计
    total = len(citations)
    unique_laws = len(groups)
    unique_articles = sum(
        len(g['articles_mentioned']) for g in groups.values()
    )

    return {
        'source_file': source_file,
        'statistics': {
            'total_citations': total,
            'unique_laws': unique_laws,
            'unique_article_references': unique_articles,
        },
        'law_groups': [
            {
                'law_name': g['law_name'],
                'original_name': g['original_name'],
                'articles_mentioned': g['articles_mentioned'],
                'citation_count': g['citation_count'],
                'citations': [
                    {
                        'raw_text': c['raw_text'],
                        'pattern_type': c['pattern_type'],
                        'article': c.get('article'),
                        'article_suffix': c.get('article_suffix'),
                        'paragraph': c.get('paragraph'),
                        'item': c.get('item'),
                        'bian': c.get('bian'),
                        'zhang': c.get('zhang'),
                        'jie': c.get('jie'),
                    }
                    for c in g['citations']
                ],
            }
            for g in groups.values()
        ],
        'all_citations': [
            {
                'raw_text': c['raw_text'],
                'pattern_type': c['pattern_type'],
                'law_name': normalize_law_name(c.get('law_name', '')),
                'article': c.get('article'),
                'article_suffix': c.get('article_suffix'),
                'paragraph': c.get('paragraph'),
                'item': c.get('item'),
                'bian': c.get('bian'),
                'zhang': c.get('zhang'),
                'jie': c.get('jie'),
                'position': c.get('position'),
            }
            for c in citations
        ],
    }


# ============================================================
# 主函数
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='法律引用批量提取器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python parse_citations.py document.docx
  python parse_citations.py contract.pdf --output citations.json
  python parse_citations.py article.txt
        """,
    )
    parser.add_argument('file', help='输入文件路径（.docx/.pdf/.txt/.md）')
    parser.add_argument('-o', '--output', help='输出 JSON 文件路径（默认输出到 stdout）')

    args = parser.parse_args()

    # 读取文件
    try:
        text = read_file(args.file)
    except (FileNotFoundError, ValueError, ImportError) as e:
        print(f"错误：{e}", file=sys.stderr)
        sys.exit(1)

    if not text.strip():
        print("警告：文件内容为空", file=sys.stderr)
        sys.exit(0)

    # 提取引用
    citations = extract_citations(text)

    if not citations:
        print("未找到法律引用", file=sys.stderr)
        result = {
            'source_file': args.file,
            'statistics': {'total_citations': 0, 'unique_laws': 0, 'unique_article_references': 0},
            'law_groups': [],
            'all_citations': [],
        }
    else:
        # 分组去重
        groups = group_by_law(citations)
        # 生成输出
        result = generate_output(citations, groups, args.file)

    # 输出
    output_json = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        Path(args.output).write_text(output_json, encoding='utf-8')
        print(f"结果已保存到：{args.output}", file=sys.stderr)
        print(f"统计：{result['statistics']['total_citations']} 条引用，"
              f"{result['statistics']['unique_laws']} 部法律，"
              f"{result['statistics']['unique_article_references']} 个条款引用", file=sys.stderr)
    else:
        print(output_json)


if __name__ == '__main__':
    main()
